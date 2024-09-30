from flask import Flask, request, jsonify, send_file
import openai
from docx import Document
import io
from io import BytesIO
import os
from flask_mysqldb import MySQL
from werkzeug.security import generate_password_hash, check_password_hash
from flask_cors import CORS
from flask_jwt_extended import JWTManager, create_access_token, jwt_required, get_jwt_identity
from googletrans import Translator
from dotenv import load_dotenv
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.oxml import OxmlElement

load_dotenv()

app = Flask(__name__)

CORS(app, resources={r"/*": {"origins": "*"}})

app.config['MYSQL_HOST'] = 'localhost'
app.config['MYSQL_USER'] = 'root'
app.config['MYSQL_PASSWORD'] = 'root123'
app.config['MYSQL_DB'] = 'login'

app.config['JWT_SECRET_KEY'] = 'mndfnfdbfdmnfvsfdvmnfdbvhdfvvbfgnbv' 
jwt = JWTManager(app)

mysql = MySQL(app)

os.environ["OPENAI_API_KEY"] = ''
openai.api_key = os.getenv("OPENAI_API_KEY")

@app.route('/register', methods=['POST'])
def register():
    try:
        data = request.json
        username = data.get('username')
        email = data.get('email')
        password = data.get('password')

        if not username or not email or not password:
            return jsonify({'message': 'All fields are required'}), 400

        cur = mysql.connection.cursor()
        
        cur.execute("SELECT * FROM users WHERE username=%s OR email=%s", (username, email))
        existing_user = cur.fetchone()

        if existing_user:
            cur.close()
            return jsonify({'message': 'Username or email already exists'}), 400

        hashed_password = generate_password_hash(password, method='pbkdf2:sha256')

        cur.execute("INSERT INTO users (username, email, password) VALUES (%s, %s, %s)", (username, email, hashed_password))
        mysql.connection.commit()
        cur.close()

        return jsonify({'message': 'User registered successfully'}), 201
    except Exception as e:
        print(f"Error: {e}")
        return jsonify({'message': 'Internal server error', 'error': str(e)}), 500
    
    
@app.route('/login', methods=['POST'])
def login():
    try:
        data = request.json
        email = data.get('email')
        password = data.get('password')

        if not email or not password:
            return jsonify({'message': 'Email and password are required'}), 400

        cur = mysql.connection.cursor()
        
        cur.execute("SELECT * FROM users WHERE email=%s", [email])
        user = cur.fetchone()
        cur.close()

        if not user:
            return jsonify({'message': 'Invalid email or password'}), 401

        user_password_hash = user[2]  # Assuming the password is the third column
        if not check_password_hash(user_password_hash, password):
            return jsonify({'message': 'Invalid email or password'}), 401

        access_token = create_access_token(identity=email)

        return jsonify({'token': access_token}), 200
    except Exception as e:
        print(f"Error: {e}")
        return jsonify({'message': 'Internal server error', 'error': str(e)}), 500

@app.route('/protected', methods=['GET'])
@jwt_required()
def protected():
    current_user = get_jwt_identity()
    return jsonify({'message': f'Hello, {current_user}'}), 200


print(f"OpenAI API Key: {os.getenv('OPENAI_API_KEY')}")

def generate_brd(summary):
    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=[
            {"role": "system", "content": """
You are tasked with generating a Business Requirements Document (BRD) based on the provided summary. 
Please ensure that the document includes all relevant sections and is formatted correctly. 
Do not include any unnecessary symbols like -,*, or /.
            """},
            {"role": "user", "content": f"Generate a detailed Business Requirement Document (BRD) based on the following summary:\n\n{summary}\n\n"
                                        "Include the following sections:\n"
                                        "1. Project Information (Project Name, Document Version, Date)\n"
                                        "2. Document History\n"
                                        "3. Table of Contents\n"
                                        "4. Introduction (Purpose, Scope, Document Overview)\n"
                                        "5. Business Objectives (Business Goal, Project Objectives, Success Criteria)\n"
                                        "6. Current Business Environment (Current Processes, Challenges, Issues)\n"
                                        "7. Proposed Solution (Description of Solution)\n"
                                        "8. Functional Requirements\n"
                                        "9. Non-Functional Requirements\n"
                                        "10. Stakeholders (List of Stakeholders, Role & Responsibilities)\n"
                                        "11. Constraints (Budgetary Constraints, Timeline, Regulatory)\n"
                                        "12. Assumptions\n"
                                        "13. Risks\n"
                                        "14. Dependencies\n"
                                        "15. Approval (Include Approval details)\n"
                                        "Format the output with sections and tables where applicable."}
        ],
        max_tokens=2000
    )
    response['choices'][0]['message']['content'] = response['choices'][0]['message']['content'].replace("*"," ")
    response['choices'][0]['message']['content'] = response['choices'][0]['message']['content'].replace("#"," ")
    return response['choices'][0]['message']['content'].strip()

def translate_text(text, target_language):
    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=[
            {"role": "system", "content": f"Translate the following text to {target_language}:"},
            {"role": "user", "content": text}
        ],
        max_tokens=1000
    )
    return response['choices'][0]['message']['content'].strip()

def create_word_document_brd(brd_text):
    document = Document()
    document.add_heading('Business Requirement Document (BRD)', 0)

    sections = brd_text.split("\n\n")
    
    for section in sections:
        if ':' in section:
            title, content = section.split(':', 1)
            
            heading = document.add_heading(title.strip(), level=1)
            run = heading.runs[0]
            run.font.size = Pt(14)
            run.font.bold = True
            
            shading_elm = parse_xml(r'<w:shd {} w:fill="D9EAD3"/>'.format(nsdecls('w')))
            heading._element.get_or_add_pPr().append(shading_elm)
            
            if "Table" in title:
                table_data = [row.strip() for row in content.strip().split('\n')]
                table = document.add_table(rows=1, cols=len(table_data[0].split('|')))
                
                hdr_cells = table.rows[0].cells
                for i, header in enumerate(table_data[0].split('|')):
                    hdr_cells[i].text = header.strip()
                    shading_elm = parse_xml(r'<w:shd {} w:fill="D9EAD3"/>'.format(nsdecls('w')))
                    hdr_cells[i]._element.get_or_add_tcPr().append(shading_elm)
                
                for row_data in table_data[1:]:
                    row_cells = table.add_row().cells
                    for i, cell in enumerate(row_data.split('|')):
                        row_cells[i].text = cell.strip()
            else:
                document.add_paragraph(content.strip())
        else:
            document.add_paragraph(section.strip())

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    
    return buffer

@app.route('/generate-brd', methods=['POST'])
def generate_brd_endpoint():
    try:
        data = request.json
        summary = data.get('summary', '')
        language = data.get('language', 'en')
        
        print(f"Summary: {summary}")
        print(f"Language: {language}")
        
        brd_content = generate_brd(summary)
        print("BRD Content generated successfully.")

        if language != 'en':
            brd_content = translate_text(brd_content, language)

        return jsonify({'brd_content': brd_content}), 200
        # buffer = create_word_document(brd_content)
        # print("Word document created successfully.")

        # return send_file(buffer, as_attachment=True, download_name="BRD.docx")
    except Exception as e:
        return jsonify({'message': 'Internal server error', 'error': str(e)}), 500

@app.route('/download-brd-doc', methods=['POST'])
def download_brd_doc():
    try:
        data = request.json
        brd_content = data.get('brd_content', '')

        buffer = create_word_document_brd(brd_content)
        return send_file(
            buffer,
            as_attachment=True,
            download_name='Business_Requirement_Document.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        return jsonify({'message': 'Internal server error', 'error': str(e)}), 500
    

def generate_frd(summary):
    try:
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": """
                You are tasked with generating a Functional Requirements Document (FRD) based on the provided summary. 
                Ensure the document includes all relevant sections and is formatted correctly. Avoid using unnecessary symbols.
                """},
                {"role": "user", "content": f"Generate a detailed Functional Requirement Document (FRD) based on the following summary:\n\n{summary}\n\n"
                                            "Include sections such as:\n"
                                            "1. Project Overview\n"
                                            "2. Scope\n"
                                            "3. Functional Requirements\n"
                                            "4. Non-Functional Requirements\n"
                                            "5. Assumptions\n"
                                            "6. Constraints\n"
                                            "7. User Interface Requirements\n"
                                            "8. Data Requirements\n"
                                            "9. Security and Compliance\n"
                                            "10. Approval Process"}
            ],
            max_tokens=2000
        )
        frd_content = response['choices'][0]['message']['content']
        frd_content = frd_content.replace("*", " ").replace("#", " ")
        print("FRD Content:", frd_content)  
        return frd_content.strip()
    except Exception as e:
        print(f"Error generating FRD from OpenAI: {e}")
        raise e 

def create_word_document_frd(frd_text):
    document = Document()
    document.add_heading('Functional Requirement Document (FRD)', 0)

    sections = frd_text.split("\n\n")
    
    for section in sections:
        if ':' in section:
            title, content = section.split(':', 1)
            
            heading = document.add_heading(title.strip(), level=1)
            run = heading.runs[0]
            run.font.size = Pt(14)
            run.font.bold = True
            
            shading_elm = parse_xml(r'<w:shd {} w:fill="D9EAD3"/>'.format(nsdecls('w')))
            heading._element.get_or_add_pPr().append(shading_elm)
            
            if "Table" in title:
                table_data = [row.strip() for row in content.strip().split('\n')]
                table = document.add_table(rows=1, cols=len(table_data[0].split('|')))
                
                hdr_cells = table.rows[0].cells
                for i, header in enumerate(table_data[0].split('|')):
                    hdr_cells[i].text = header.strip()
                    shading_elm = parse_xml(r'<w:shd {} w:fill="D9EAD3"/>'.format(nsdecls('w')))
                    hdr_cells[i]._element.get_or_add_tcPr().append(shading_elm)
                
                for row_data in table_data[1:]:
                    row_cells = table.add_row().cells
                    for i, cell in enumerate(row_data.split('|')):
                        row_cells[i].text = cell.strip()
            else:
                document.add_paragraph(content.strip())
        else:
            document.add_paragraph(section.strip())

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    
    return buffer

@app.route('/generate-frd', methods=['POST'])
def generate_frd_endpoint():
    try:
        data = request.json
        summary = data.get('summary', '')
        language = data.get('language', 'en')
        
        print(f"Received summary: {summary}")
        print(f"Requested language: {language}")
        
        # Generate FRD content
        frd_content = generate_frd(summary)
        print("FRD content generated successfully.")

        # Translate if needed
        if language != 'en':
            frd_content = translate_text(frd_content, language)
            print("FRD content translated successfully.")
        
        return jsonify({'frd_content': frd_content}), 200
    except Exception as e:
        print(f"Error generating FRD: {e}")
        return jsonify({'message': 'Internal server error', 'error': str(e)}), 500
 
@app.route('/download-frd-doc', methods=['POST'])
def download_frd_doc():
    try:
        data = request.json
        frd_content = data.get('frd_content', '')

        buffer = create_word_document_frd(frd_content)
        return send_file(
            buffer,
            as_attachment=True,
            download_name='Functional_Requirement_Document.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        return jsonify({'message': 'Internal server error', 'error': str(e)}), 500
 
    
def generate_prd(summary):
    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
       messages=[
            {"role": "system", "content": """
You are tasked with generating a Product Requirements Document (PRD) based on the provided summary. 
Ensure the document includes all relevant sections, such as Executive Summary, Target Audience, Product Vision and Mission, Features and Functionality, Non-Functional Requirements, and Technical Requirements.
Do not use unnecessary symbols like -,*, or /.
            """},
            {"role": "user", "content": f"Generate a detailed Product Requirement Document (PRD) based on the following summary:\n\n{summary}\n\n"
                                        "Ensure that the document includes the following sections:\n"
                                        "1. Executive Summary\n"
                                        "2. Target Audience\n"
                                        "3. Product Vision and Mission\n"
                                        "4. Features and Functionality\n"
                                        "5. Non-Functional Requirements\n"
                                        "6. Technical Requirements\n"
                                        "7. Timeline and Milestones\n"
                                        "8. Success Metrics\n"
                                        "9. Appendices\n"}
        ],
        max_tokens=2000
    )
    response['choices'][0]['message']['content'] = response['choices'][0]['message']['content'].replace("*"," ")
    response['choices'][0]['message']['content'] = response['choices'][0]['message']['content'].replace("#"," ")
    prd_content = response['choices'][0]['message']['content']
    return prd_content.strip()
    
def create_word_document_prd(prd_text):
    document = Document()
    document.add_heading('Product Requirement Document (PRD)', 0)

    sections = prd_text.split("\n\n")
    
    for section in sections:
        if ':' in section:
            title, content = section.split(':', 1)
            
            heading = document.add_heading(title.strip(), level=1)
            run = heading.runs[0]
            run.font.size = Pt(14)
            run.font.bold = True
            
            shading_elm = parse_xml(r'<w:shd {} w:fill="D9EAD3"/>'.format(nsdecls('w')))
            heading._element.get_or_add_pPr().append(shading_elm)
            
            if "Table" in title:
                table_data = [row.strip() for row in content.strip().split('\n')]
                table = document.add_table(rows=1, cols=len(table_data[0].split('|')))
                
                hdr_cells = table.rows[0].cells
                for i, header in enumerate(table_data[0].split('|')):
                    hdr_cells[i].text = header.strip()
                    shading_elm = parse_xml(r'<w:shd {} w:fill="D9EAD3"/>'.format(nsdecls('w')))
                    hdr_cells[i]._element.get_or_add_tcPr().append(shading_elm)
                
                for row_data in table_data[1:]:
                    row_cells = table.add_row().cells
                    for i, cell in enumerate(row_data.split('|')):
                        row_cells[i].text = cell.strip()
            else:
                document.add_paragraph(content.strip())
        else:
            document.add_paragraph(section.strip())

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    
    return buffer

@app.route('/generate-prd', methods=['POST'])
def generate_prd_endpoint():
    try:
        data = request.json
        summary = data.get('summary', '')
        language = data.get('language', 'en')
        
        print(f"Summary: {summary}")
        print(f"Language: {language}")
        
        prd_content = generate_prd(summary)
        print("PRD Content generated successfully.")

        if language != 'en':
            prd_content = translate_text(prd_content, language)

        return jsonify({'prd_content': prd_content}), 200
        # buffer = create_word_document(prd_content)
        # print("Word document created successfully.")
        # return send_file(buffer, as_attachment=True, download_name="PRD.docx", mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    except Exception as e:
        return jsonify({'message': 'Internal server error', 'error': str(e)}), 500

@app.route('/download-prd-doc', methods=['POST'])
def download_prd_doc():
    try:
        data = request.json
        prd_content = data.get('prd_content', '')

        buffer = create_word_document_prd(prd_content)
        return send_file(
            buffer,
            as_attachment=True,
            download_name='Product_Requirement_Document.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        return jsonify({'message': 'Internal server error', 'error': str(e)}), 500
    
    
def generate_technical_documentation(summary):
    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=[
            {"role": "system", "content": """
Your are technical writers, project managers, members of a development team or experts on the product or service,
and you are responsible for creating technical documentation for a product or service. 
Your task is to generate comprehensive technical documentation.

The technical documentation should include the following sections:

- Introduction
- Title page
- Table of contents
- Summary
- Conclusions
- List of figures
- Project plans
- Results
- User guides
- Business proposals
- Header
- Product requirements
- Appendix
- Materials
- Work
- Body
- Code Documentation
- Procedure
            """},
            {"role": "user", "content": f"""Generate a detailed Technical Document based on the following summary:\n\n{summary}\n\n
                                        Ensure that the document includes all the sections mentioned above, 
                                        and do not include any -,*, or / in the document.
                                        """}
        ],
        max_tokens=2000
    )
    response['choices'][0]['message']['content'] = response['choices'][0]['message']['content'].replace("*"," ")
    response['choices'][0]['message']['content'] = response['choices'][0]['message']['content'].replace("#"," ")
    return response['choices'][0]['message']['content'].strip()


def create_word_document_techdoc(tech_doc_content):
    document = Document()
    document.add_heading("Technical Document", 0)

    sections = tech_doc_content.split("\n\n")
    
    for section in sections:
        if ':' in section:
            title, content = section.split(':', 1)
            
            heading = document.add_heading(title.strip(), level=1)
            run = heading.runs[0]
            run.font.size = Pt(14)
            run.font.bold = True
            
            shading_elm = parse_xml(r'<w:shd {} w:fill="D9EAD3"/>'.format(nsdecls('w')))
            heading._element.get_or_add_pPr().append(shading_elm)
            
            if "Table" in title:
                table_data = [row.strip() for row in content.strip().split('\n')]
                table = document.add_table(rows=1, cols=len(table_data[0].split('|')))
                
                hdr_cells = table.rows[0].cells
                for i, header in enumerate(table_data[0].split('|')):
                    hdr_cells[i].text = header.strip()
                    shading_elm = parse_xml(r'<w:shd {} w:fill="D9EAD3"/>'.format(nsdecls('w')))
                    hdr_cells[i]._element.get_or_add_tcPr().append(shading_elm)
                
                for row_data in table_data[1:]:
                    row_cells = table.add_row().cells
                    for i, cell in enumerate(row_data.split('|')):
                        row_cells[i].text = cell.strip()
            else:
                document.add_paragraph(content.strip())
        else:
            document.add_paragraph(section.strip())

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    
    return buffer
    
@app.route('/generate-technical-document', methods=['POST'])
def generate_technical_document_endpoint():
    try:
        data = request.json
        summary = data.get('summary', '')
        language = data.get('language', 'en')

        tech_doc_content = generate_technical_documentation(summary)
        print("Technical Document Content generated successfully.")

        if language != 'en':
            tech_doc_content = translate_text(tech_doc_content, language)

        return jsonify({'tech_doc_content': tech_doc_content}), 200
    except Exception as e:
        return jsonify({'message': 'Internal server error', 'error': str(e)}), 500

@app.route('/download-tech-doc', methods=['POST'])
def download_tech_doc():
    try:
        data = request.json
        tech_doc_content = data.get('tech_doc_content', '')

        buffer = create_word_document_techdoc(tech_doc_content)
        return send_file(
            buffer,
            as_attachment=True,
            download_name='Technical_Document.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        return jsonify({'message': 'Internal server error', 'error': str(e)}), 500

def generate_user_manual(summary):
    try:
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": """
                You are tasked with generating a User Manual based on the provided summary. 
                Ensure the document includes all relevant sections and is formatted correctly. Avoid using unnecessary symbols.
                """},
                {"role": "user", "content": f"Generate a detailed User Manual based on the following summary:\n\n{summary}\n\n"
                                            "Include sections such as:\n"
                                            "1. Introduction (Overview, Purpose, Scope)\n"
                                            "2. Getting Started (Installation, Setup, Configuration)\n"
                                            "3. User Interface Overview (Description of UI components)\n"
                                            "4. Features and Functionality (List of features and instructions)\n"
                                            "5. Troubleshooting (Common issues and fixes)\n"
                                            "6. FAQs (Frequently asked questions)\n"
                                            "7. Support (How to get help)"}
            ],
            max_tokens=2000
        )
        user_manual_content = response['choices'][0]['message']['content']
        user_manual_content = user_manual_content.replace("*", " ").replace("#", " ")
        print("User Manual Content:", user_manual_content)
        return user_manual_content.strip()
    except Exception as e:
        print(f"Error generating User Manual from OpenAI: {e}")
        raise e

def create_word_document_usermanual(usermanual_text):
    document = Document()
    document.add_heading('User Manual', 0)

    sections = usermanual_text.split("\n\n")
    
    for section in sections:
        if ':' in section:
            title, content = section.split(':', 1)
            
            heading = document.add_heading(title.strip(), level=1)
            run = heading.runs[0]
            run.font.size = Pt(14)
            run.font.bold = True
            
            shading_elm = parse_xml(r'<w:shd {} w:fill="D9EAD3"/>'.format(nsdecls('w')))
            heading._element.get_or_add_pPr().append(shading_elm)
            
            if "Table" in title:
                table_data = [row.strip() for row in content.strip().split('\n')]
                table = document.add_table(rows=1, cols=len(table_data[0].split('|')))
                
                hdr_cells = table.rows[0].cells
                for i, header in enumerate(table_data[0].split('|')):
                    hdr_cells[i].text = header.strip()
                    shading_elm = parse_xml(r'<w:shd {} w:fill="D9EAD3"/>'.format(nsdecls('w')))
                    hdr_cells[i]._element.get_or_add_tcPr().append(shading_elm)
                
                for row_data in table_data[1:]:
                    row_cells = table.add_row().cells
                    for i, cell in enumerate(row_data.split('|')):
                        row_cells[i].text = cell.strip()
            else:
                document.add_paragraph(content.strip())
        else:
            document.add_paragraph(section.strip())

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    
    return buffer

@app.route('/generate-user-manual', methods=['POST'])
def generate_user_manual_endpoint():
    try:
        data = request.json
        summary = data.get('summary', '')
        language = data.get('language', 'en')
        
        print(f"Received summary: {summary}")
        print(f"Requested language: {language}")
        
        # Generate User Manual content
        user_manual_content = generate_user_manual(summary)
        print("User Manual content generated successfully.")

        # Translate if needed
        if language != 'en':
            user_manual_content = translate_text(user_manual_content, language)
            print("User Manual content translated successfully.")
        
        return jsonify({'user_manual_content': user_manual_content}), 200
    except Exception as e:
        print(f"Error generating User Manual: {e}")
        return jsonify({'message': 'Internal server error', 'error': str(e)}), 500


@app.route('/download-usermanual-doc', methods=['POST'])
def download_usermanual_doc():
    try:
        data = request.json
        user_manual_content = data.get('user_manual_content', '')

        buffer = create_word_document_usermanual(user_manual_content)
        return send_file(
            buffer,
            as_attachment=True,
            download_name='User_Manual.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        return jsonify({'message': 'Internal server error', 'error': str(e)}), 500

def generate_qrd(summary):
    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=[
            {
                "role": "system",
                "content": ("""
                You are an expert in creating only Quality Requirements Documents.
                A QRD describes the customer's expectations for the final product's quality. It comprises different metrics, factors, and criteria. Quality requirements documents might revolve around customer experience, maintainability, usability, availability, consistency, and reliability. 
                
                Important: Make sure the document you generate is titled 'Quality Requirements Document' and reflects all sections related to the QRD. Ensure that no other document name is mentioned, and do not include any symbols such as -, *, or / in the document.
                """)
            },
            {"role": "user", "content": f"""Generate a detailed Quality Requirements Document based on the following summary:\n\n{summary}\n\nEnsure that the document is correctly titled as 'Quality Requirements Document'."""}
        ],
        max_tokens=2000
    )
    response['choices'][0]['message']['content'] = response['choices'][0]['message']['content'].replace("*"," ")
    response['choices'][0]['message']['content'] = response['choices'][0]['message']['content'].replace("#"," ")
    return response['choices'][0]['message']['content'].strip()


def create_word_document_qrd(qrd_text):
    document = Document()
    document.add_heading('Quality Requirements Document(QRD)', 0)

    sections = qrd_text.split("\n\n")
    
    for section in sections:
        if ':' in section:
            title, content = section.split(':', 1)
            
            heading = document.add_heading(title.strip(), level=1)
            run = heading.runs[0]
            run.font.size = Pt(14)
            run.font.bold = True
            
            shading_elm = parse_xml(r'<w:shd {} w:fill="D9EAD3"/>'.format(nsdecls('w')))
            heading._element.get_or_add_pPr().append(shading_elm)
            
            if "Table" in title:
                table_data = [row.strip() for row in content.strip().split('\n')]
                table = document.add_table(rows=1, cols=len(table_data[0].split('|')))
                
                hdr_cells = table.rows[0].cells
                for i, header in enumerate(table_data[0].split('|')):
                    hdr_cells[i].text = header.strip()
                    shading_elm = parse_xml(r'<w:shd {} w:fill="D9EAD3"/>'.format(nsdecls('w')))
                    hdr_cells[i]._element.get_or_add_tcPr().append(shading_elm)
                
                for row_data in table_data[1:]:
                    row_cells = table.add_row().cells
                    for i, cell in enumerate(row_data.split('|')):
                        row_cells[i].text = cell.strip()
            else:
                document.add_paragraph(content.strip())
        else:
            document.add_paragraph(section.strip())

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    
    return buffer

@app.route('/generate-qrd', methods=['POST'])
def generate_qrd_endpoint():
    try:
        data = request.json
        summary = data.get('summary', '')
        language = data.get('language', 'en')

        qrd_content = generate_qrd(summary)
        print("QRD content generated successfully.")

        if language != 'en':
            qrd_content = translate_text(qrd_content, language)

        return jsonify({'qrd_content': qrd_content}), 200
    except Exception as e:
        return jsonify({'message': 'Internal server error', 'error': str(e)}), 500

@app.route('/download-qrd-doc', methods=['POST'])
def download_qrd_doc():
    try:
        data = request.json
        qrd_content = data.get('qrd_content', '')

        buffer = create_word_document_qrd(qrd_content)
        return send_file(
            buffer,
            as_attachment=True,
            download_name='Quality_Requirement_Document.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        return jsonify({'message': 'Internal server error', 'error': str(e)}), 500

def generate_mrd(summary):
    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=[
            {"role": "system", "content": """
                You are an expert in creating Market Requirements Documents (MRD).
                An MRD defines the product's market requirements, including target market, market needs, competitive landscape, key product features, pricing strategy, and expected sales. 
                Ensure that the document includes all necessary sections and do not use any symbols like -, *, or / in the document.
            """},
            {"role": "user", "content": f"""Generate a detailed Market Requirements Document based on the following summary:\n\n{summary}\n\nEnsure that the document includes all the sections mentioned above, and do not include any symbols such as -, *, or / in the document."""}
        ],
        max_tokens=2000
    )
    response['choices'][0]['message']['content'] = response['choices'][0]['message']['content'].replace("*"," ")
    response['choices'][0]['message']['content'] = response['choices'][0]['message']['content'].replace("#"," ")
    return response['choices'][0]['message']['content'].strip()

def create_word_document_mrd(mrd_text):
    document = Document()
    document.add_heading("Market Requirements Document", 0)

    sections = mrd_text.split("\n\n")
    
    for section in sections:
        if ':' in section:
            title, content = section.split(':', 1)
            
            heading = document.add_heading(title.strip(), level=1)
            run = heading.runs[0]
            run.font.size = Pt(14)
            run.font.bold = True
            
            shading_elm = parse_xml(r'<w:shd {} w:fill="D9EAD3"/>'.format(nsdecls('w')))
            heading._element.get_or_add_pPr().append(shading_elm)
            
            if "Table" in title:
                table_data = [row.strip() for row in content.strip().split('\n')]
                table = document.add_table(rows=1, cols=len(table_data[0].split('|')))
                
                hdr_cells = table.rows[0].cells
                for i, header in enumerate(table_data[0].split('|')):
                    hdr_cells[i].text = header.strip()
                    shading_elm = parse_xml(r'<w:shd {} w:fill="D9EAD3"/>'.format(nsdecls('w')))
                    hdr_cells[i]._element.get_or_add_tcPr().append(shading_elm)
                
                for row_data in table_data[1:]:
                    row_cells = table.add_row().cells
                    for i, cell in enumerate(row_data.split('|')):
                        row_cells[i].text = cell.strip()
            else:
                document.add_paragraph(content.strip())
        else:
            document.add_paragraph(section.strip())

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    
    return buffer

@app.route('/generate-mrd', methods=['POST'])
def generate_mrd_endpoint():
    try:
        data = request.json
        summary = data.get('summary', '')
        language = data.get('language', 'en')

        mrd_content = generate_mrd(summary)
        print("MRD content generated successfully.")

        if language != 'en':
            mrd_content = translate_text(mrd_content, language)

        return jsonify({'mrd_content': mrd_content}), 200
    except Exception as e:
        return jsonify({'message': 'Internal server error', 'error': str(e)}), 500

@app.route('/download-mrd-doc', methods=['POST'])
def download_mrd_doc():
    try:
        data = request.json
        mrd_content = data.get('mrd_content', '')

        buffer = create_word_document_mrd(mrd_content)
        return send_file(
            buffer,
            as_attachment=True,
            download_name='Market_Requirement_Document.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        return jsonify({'message': 'Internal server error', 'error': str(e)}), 500

def generate_tpd(summary):
    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=[
            {"role": "system", "content": """
                You are a QA manager or test lead responsible for writing a test plan document. 
                The document should be well-organized and professional, including sections such as Project Information, Test Strategy, Test Scope, Test Plan, Test Design, Test Execution, Test Closure, and Appendices. 
                Ensure that the document includes all necessary sections and avoid using any symbols like -, *, or / in the document.
            """},
            {"role": "user", "content": f"""Generate a detailed Test Plan Document based on the following summary:\n\n{summary}\n\nEnsure that the document includes all the sections mentioned above, and do not include any symbols such as -, *, or / in the document."""}
        ],
        max_tokens=2000
    )
    response['choices'][0]['message']['content'] = response['choices'][0]['message']['content'].replace("*"," ")
    response['choices'][0]['message']['content'] = response['choices'][0]['message']['content'].replace("#"," ")
    return response['choices'][0]['message']['content'].strip()

def create_word_document_tpd(tpd_text):
    document = Document()
    document.add_heading('Test Plan Document (TPD)', 0)

    sections = tpd_text.split("\n\n")
    
    for section in sections:
        if ':' in section:
            title, content = section.split(':', 1)
            
            heading = document.add_heading(title.strip(), level=1)
            run = heading.runs[0]
            run.font.size = Pt(14)
            run.font.bold = True
            
            shading_elm = parse_xml(r'<w:shd {} w:fill="D9EAD3"/>'.format(nsdecls('w')))
            heading._element.get_or_add_pPr().append(shading_elm)
            
            if "Table" in title:
                table_data = [row.strip() for row in content.strip().split('\n')]
                table = document.add_table(rows=1, cols=len(table_data[0].split('|')))
                
                hdr_cells = table.rows[0].cells
                for i, header in enumerate(table_data[0].split('|')):
                    hdr_cells[i].text = header.strip()
                    shading_elm = parse_xml(r'<w:shd {} w:fill="D9EAD3"/>'.format(nsdecls('w')))
                    hdr_cells[i]._element.get_or_add_tcPr().append(shading_elm)
                
                for row_data in table_data[1:]:
                    row_cells = table.add_row().cells
                    for i, cell in enumerate(row_data.split('|')):
                        row_cells[i].text = cell.strip()
            else:
                document.add_paragraph(content.strip())
        else:
            document.add_paragraph(section.strip())

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    
    return buffer

@app.route('/generate-tpd', methods=['POST'])
def generate_tpd_endpoint():
    try:
        data = request.json
        summary = data.get('summary', '')
        language = data.get('language', 'en')

        tpd_content = generate_tpd(summary)
        print("TPD content generated successfully.")

        if language != 'en':
            tpd_content = translate_text(tpd_content, language)

        return jsonify({'tpd_content': tpd_content}), 200
    except Exception as e:
        return jsonify({'message': 'Internal server error', 'error': str(e)}), 500

@app.route('/download-tpd-doc', methods=['POST'])
def download_tpd_doc():
    try:
        data = request.json
        tpd_content = data.get('tpd_content', '')

        buffer = create_word_document_tpd(tpd_content)
        return send_file(
            buffer,
            as_attachment=True,
            download_name='Test_Plan_Document.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        return jsonify({'message': 'Internal server error', 'error': str(e)}), 500

def generate_api_documentation(summary):
    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=[
            {
                "role": "system",
                "content": """
                You are an API Developer, Technical Writer, Product Manager, and API Evangelist.
                Your task is to generate a comprehensive API documentation based on the provided summary.
                Include the following sections: Introduction, Authentication, Endpoints, Data Structures, Rate Limiting, Caching, Changelog, Examples, and Contact Information.
                """
            },
            {"role": "user", "content": f"Generate a detailed API documentation based on the following summary:\n\n{summary}\n\n"}
        ]
    )
    response['choices'][0]['message']['content'] = response['choices'][0]['message']['content'].replace("*"," ")
    response['choices'][0]['message']['content'] = response['choices'][0]['message']['content'].replace("#"," ")
    return response['choices'][0]['message']['content'].strip()

def create_word_document_api(api_content):
    document = Document()
    document.add_heading("API Documentation", 0)

    sections = api_content.split("\n\n")
    
    for section in sections:
        if ':' in section:
            title, content = section.split(':', 1)
            
            heading = document.add_heading(title.strip(), level=1)
            run = heading.runs[0]
            run.font.size = Pt(14)
            run.font.bold = True
            
            shading_elm = parse_xml(r'<w:shd {} w:fill="D9EAD3"/>'.format(nsdecls('w')))
            heading._element.get_or_add_pPr().append(shading_elm)
            
            if "Table" in title:
                table_data = [row.strip() for row in content.strip().split('\n')]
                table = document.add_table(rows=1, cols=len(table_data[0].split('|')))
                
                hdr_cells = table.rows[0].cells
                for i, header in enumerate(table_data[0].split('|')):
                    hdr_cells[i].text = header.strip()
                    shading_elm = parse_xml(r'<w:shd {} w:fill="D9EAD3"/>'.format(nsdecls('w')))
                    hdr_cells[i]._element.get_or_add_tcPr().append(shading_elm)
                
                for row_data in table_data[1:]:
                    row_cells = table.add_row().cells
                    for i, cell in enumerate(row_data.split('|')):
                        row_cells[i].text = cell.strip()
            else:
                document.add_paragraph(content.strip())
        else:
            document.add_paragraph(section.strip())

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    
    return buffer

@app.route('/generate-api-doc', methods=['POST'])
def generate_document():
    try:
        data = request.json
        summary = data.get("summary")
        language = data.get("language")

        api_content = generate_api_documentation(summary)
        print("API document created")
        print(api_content)

        if language != 'English':
            api_content = translate_text(api_content, language)

        return jsonify({'api_content': api_content})
    except Exception as e:
        return jsonify({'message': 'Internal server error', 'error': str(e)}), 500


@app.route('/download-api-doc', methods=['POST'])
def download_api_doc():
    try:
        data = request.json
        api_content = data.get('api_content', '')

        buffer = create_word_document_api(api_content)
        return send_file(
            buffer,
            as_attachment=True,
            download_name='API_Document.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        return jsonify({'message': 'Internal server error', 'error': str(e)}), 500

if __name__ == '__main__':
    app.run(debug=True)